from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.firefox.options import Options
import re
from bs4 import BeautifulSoup as bs
import time
from docx import Document
from docx.shared import Inches
import glob
import os
import requests
from docxcompose.composer import Composer
from datetime import datetime as dt
import sys

class Willhaben():

    driver = ""

    def call_url(self):
        """
        Function loads Firefox driver and open willhaben.at
        return True if page is loaded
        """
        try:
            #If you don't have X server headless mode activates
            options = Options()
            options.add_argument('-headless')
            self.driver = webdriver.Firefox(executable_path='geckodriver', options=options)
            #Load Firefox driver
            #self.driver = webdriver.Firefox()
            #Open willhaben.at
            self.driver.get("https://www.willhaben.at/iad/kaufen-und-verkaufen/wohnen-haushalt-gastronomie")
            WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.XPATH, '//*[text()="Cookies akzeptieren"]'))).click()
            #Check the title correct page is loaded
            assert "Wohnen" in self.driver.title
            return True
        except AssertionError:
            print("\tMaybe wrong page was loaded!")
        except Exception as e:
            print("\tError occured!<call_url>")
            print("\t",str(e))


    def filter_items(self):
        """
        Function filters free stuff on willhaben.at in House category
        return filtered url name
        """
        try:
            #Look for search input field, clear and write text and submit
            elem = self.driver.find_element_by_id('autocomplete-input')
            elem.clear()
            elem.send_keys("kostenloser")
            elem.send_keys(Keys.RETURN)
            assert "Deine Suche hat keine Treffer erzielt." not in self.driver.page_source
                #Click Preis option and select kostenlose value
            if WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, 'Preis'))):
                self.driver.find_element_by_id('Preis').click()
                self.driver.find_element_by_partial_link_text('zu verschenken').click()
                time.sleep(9)
            return re.search(r'app-argument=(.*?)"', self.driver.page_source).group(1).replace('&amp;', '&')
        except AssertionError:
            print("\tThere is no any result! Exit.")
        except Exception as e:
            print("Error occured!<filter_items>")
            print("\t",str(e))

    def select_results(self, wurl):
        """
        Function parses first 25 result
        return a list with url names
        """
        try:
            url_list = []
            #Invoke url which is consist of filtered result
            self.driver.get(wurl)
            time.sleep(5)
            whtml = self.driver.page_source
            soup = bs(whtml, 'html.parser')
            #Gain urls from filtered results
            for i in soup.find_all('a', itemprop='url'):
                url_list.append("".join(['https://willhaben.at', i.get('href')]))
            return url_list
        except Exception as e:
            print("\tError occured!<select_result>")
            print("\t",str(e))


    def collect_results(self, wurl):
        """
        Function gain useful data from advertisement
        return a list with consists of advertisement data and pictures of urls
        """
        try:
            result = []
            self.driver.get(wurl)
            time.sleep(5)
            whtml = self.driver.page_source
            soup = bs(whtml, 'html.parser')
            #First item consists of filename for docx
            result.append(wurl.split('/')[-2])
            #Image urls
            for i in soup.find_all('img', class_="image"):
                result.append(i.get('src'))
            #Default image url
            for i in soup.find_all("a", class_="img-link-no-script"):
                result.append(i.get('src'))
            #Willhabe code
            result.append([i.get_text() for i in soup.select('#advert-info-whCode')][0])
            #Adv created
            result.append([i.get_text() for i in soup.select('#advert-info-dateTime')][0])
            #Adv description
            for i in zip(soup.find_all('span',class_="col-2-desc"), soup.find_all('div',class_="col-2-body")):
                result.append("".join([str(i[0].get_text()).strip(), str(i[1].get_text()).strip()]))
            #Detailed description
            result.append(str(soup.find('div', class_='description').get_text()).strip())
            #Seller data
            seller_data = eval(soup.body['data-tealium'])
            result.append(seller_data['tmsData']['seller_name'])
            result.append(seller_data['tmsData']['price'])
            result.append(seller_data['tmsData']['post_code'])
            result.append(seller_data['tmsData']['region_level_2'].replace('%C3%B6', 'ö').replace('%2C', ''))
            result.append(seller_data['tmsData']['region_level_3'].replace('%C3%B6', 'ö').replace('%2C', ''))
            return result
        except Exception as e:
            print("\tError occured!<collect_result>")
            print("\t", str(e))


    def download_jpg(self, wurl):
        """
        Function downloads image from specific url, filename is the end of the url
        return True if downloaded succesfully
        """
        try:
            #Downloads image and writes to file
            with open('willhaben_img_' + wurl.split('/')[-1], 'wb') as f:
                f.write(requests.get(wurl).content)
            return True
        except Exception as e:
            print("\tError occured!<download_jpg>")
            print("\t", str(e))


    def add_adv(self, text):
        """
        Function writes advertisement and its images to docx file
        """
        try:
            #Opens docx document object        
            document = Document()
            for i in text:
                #Searches urls then downloads them
                if i.startswith('https'):
                    if self.download_jpg(i):
                        document.add_picture('willhaben_img_' + i.split('/')[-1], width=Inches(2.5))
                else:
                    document.add_paragraph(i)
            document.add_page_break()
            document.save('willhaben_' + text[0][:20] + '.docx')
            #Cleans up   
            for i in glob.glob('willhaben_img_*.jpg'):
                os.remove(i)
            return True
        except Exception as e:
            print("\tError occured!<add_adv>")
            print("\t", str(e))


    def merge_docx(self):
        """
        Function merges multiple docx documents
        """
        try:
            #Creates composer object
            composer = Composer(Document())
            for i in glob.glob('willhaben*.docx'):
                tmp = Document(i)
                composer.append(tmp)
                base_path = os.path.dirname(os.path.abspath(__file__))
                out_file = 'out/kostenloser_produkten_' + dt.strftime(dt.now(), "%Y-%m-%dT%H-%M-%S") + '.docx'
            composer.save("/".join([base_path, out_file]))
            #Cleans up    
            for i in glob.glob('willhaben*.docx'):
                os.remove(i)
            return True
        except Exception as e:
            print("\tError occured!<merge_docx>")
            print("\t", str(e))


    def __del__(self):
        self.driver.quit()


def main():
    try:
        #Creates instance
        session1 = Willhaben()
        #Opens firefox and invokes willhaben.at 
        if session1.call_url():
	#Searches and filters based on criteria
            wurl = session1.filter_items()
            if wurl :
                #list contains url of first 25 hits
                result_list = session1.select_results(wurl)
                if len(result_list) != 0:
                    for i in result_list:
                        #Downloads and creates docx for advertisements
                        session1.add_adv(session1.collect_results(i))
                    session1.merge_docx()
                else:
                    print('\tThere is no any result to generate riport!')
            else:
                print('\tThere is no any free stuff on willhaben.at?!Peculiar...')
    except Exception as e:
        print("\tError occured!")
        print('\t', str(e))
        sys.exit(1)

if __name__ == '__main__':
    main()
